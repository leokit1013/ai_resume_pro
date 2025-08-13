from import_python_packages import *

st.set_page_config(page_title="Enhance Existing Resume", layout="wide")

if "token" not in st.session_state:
    st.switch_page("login.py")

res = requests.post(f"{BACKEND_URL}/validate-token", json={"token": st.session_state["token"]})
if res.status_code != 200:
    st.switch_page("login.py")
    
# hide navbar and footer
hide_st_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    [data-testid="stSidebarNav"] {display: none;}
    </style>
"""
st.markdown(hide_st_style, unsafe_allow_html=True)

if "email" not in st.session_state:
    st.warning("Please login to access this page.")
    st.switch_page("app.py")

if st.session_state["usage_count"] >= 2 and not st.session_state["subscribed"]:
    st.warning("You've exceeded your free limit.")
    if st.button("Go to Payment Page"):
        st.switch_page("pages/payment_page.py")
    st.stop()

# Call update once only
if "used_this_tool" not in st.session_state:
    update_usage(st.session_state["email"])
    st.session_state["usage_count"] += 1
    st.session_state["used_this_tool"] = True

# st.title("Resume Builder Tool")





import os
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
from dotenv import load_dotenv
from langdetect import detect
import easyocr
import numpy as np
from io import BytesIO
from fpdf import FPDF
from docx import Document
from fpdf import FPDF
import unicodedata
import re


load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("API_KEY not found in .env file or environment variables.")

genai.configure(api_key=GOOGLE_API_KEY)

model = genai.GenerativeModel(
    'gemini-1.5-flash-8b-001',
    generation_config={
        "temperature": 0.7,
        "top_p": 0.95,
        "top_k": 40,
        "max_output_tokens": 8192,
        "response_mime_type": "text/plain",
    }
)

def detect_language(text):
    return detect(text)

def read_image(file, lang):
    image = Image.open(file)
    image_np = np.array(image)
    reader = easyocr.Reader([lang, 'en'], gpu=False)
    result = reader.readtext(image_np, detail=0)
    return ' '.join(result)




def clean_text(text):
    # Normalize Unicode and remove unsupported characters for FPDF
    return unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")



# ---------- Markdown to HTML Converter ----------
def markdown_to_html(text):
    html = text
    html = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", html)  # Bold
    html = re.sub(r"\*(.*?)\*", r"<em>\1</em>", html)              # Italic
    html = html.replace("\n", "<br>")                              # Newlines to <br>
    return f"<div style='font-family:Arial; font-size:15px;'>{html}</div>"

# ---------- Save as DOCX ----------
def save_as_docx(text, filename="resume.docx"):
    doc = Document()

    lines = text.split("\n")
    for line in lines:
        paragraph = doc.add_paragraph()

        # Bold: **text**
        line = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", line)
        # Italic: *text*
        line = re.sub(r"\*(.*?)\*", r"<i>\1</i>", line)

        # Split line by HTML tags for bold and italic
        tokens = re.split(r"(<b>.*?</b>|<i>.*?</i>)", line)
        for token in tokens:
            if token.startswith("<b>") and token.endswith("</b>"):
                run = paragraph.add_run(token[3:-4])
                run.bold = True
            elif token.startswith("<i>") and token.endswith("</i>"):
                run = paragraph.add_run(token[3:-4])
                run.italic = True
            else:
                paragraph.add_run(token)
    doc.save(filename)
    return filename





def extract_text_from_file(file):
    try:
 
        if file.name.endswith(".txt"):
            return file.read().decode("utf-8")

        elif file.name.endswith(".docx"):
            doc = Document(file)
            return "\n".join(p.text for p in doc.paragraphs)

        elif file.name.endswith((".jpg", ".jpeg", ".png")):
            temp_text = read_image(file, 'en')
            lang = detect_language(temp_text)
            return read_image(file, lang)

        else:
            st.error("Unsupported file type.")
            return None

    except Exception as e:
        st.error(f"Text extraction failed: {e}")
        return None

def summarize_with_gemini(text, prompt_template):
    try:
        prompt = prompt_template.format(text=text)
        response = model.generate_content(prompt)
        return response.text if response.text else None
    except Exception as e:
        st.error(f" API Error: {e}")
        return None


def main():
    
    st.title("🧠 Enhance Existing Resume")

    st.subheader("Upload your existing resume and get it improved and more impactful")
    file = st.file_uploader("Upload Resume (PDF, DOCX, TXT, Image)", type=["pdf", "docx", "txt", "jpg", "jpeg", "png"])

    improve_prompt_template = (
        "You are an expert resume editor. Improve and rewrite the resume below to make it more impactful, clean, and suitable for top jobs.\n\n"
        "Resume:\n{text}\n\n"
        "Return the improved resume as plain text."
    )

    if file and st.button("🚀 Improve Resume"):
        with st.spinner("Extracting and improving resume..."):
            raw_text = extract_text_from_file(file)
            if raw_text:
                improved = summarize_with_gemini(raw_text, improve_prompt_template)
                if improved:
                    st.success("✅ Improved Resume")
                    st.markdown(markdown_to_html(improved), unsafe_allow_html=True)

                    
                    docx_path = save_as_docx(improved, "improved_resume.docx")
                    # pdf_path = save_as_pdf(improved, "improved_resume.pdf")

                    st.download_button("📄 Download as DOCX", data=open(docx_path, "rb"), file_name="improved_resume.docx")
                    st.download_button("📄 Download as TXT", data=improved, file_name="improved_resume.txt")

                else:
                    st.error("Improvement failed.")
            else:
                st.error("Could not extract text from the uploaded file.")

if __name__ == "__main__":
    main()
